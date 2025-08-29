# app.py — Anki Deck Creator (Manual-only, polished UX, robust)
# Dependencies:
# streamlit openai genanki gTTS pypdf python-docx unidecode requests beautifulsoup4 readability-lxml
# scikit-learn pandas pytube youtube-transcript-api

from __future__ import annotations

import os, io, re, json, time, random, tempfile, hashlib
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st
from openai import OpenAI

import genanki
from gtts import gTTS
from gtts.lang import tts_langs
from unidecode import unidecode

from pypdf import PdfReader
import docx as docxlib

# Optional deps
try:
    import requests
    from bs4 import BeautifulSoup
    try:
        from readability import Document
    except Exception:
        Document = None
except Exception:
    requests = None
    BeautifulSoup = None
    Document = None

try:
    from pytube import YouTube
    from youtube_transcript_api import YouTubeTranscriptApi
    YOUTUBE_DEPS_OK = True
except Exception:
    YOUTUBE_DEPS_OK = False
    YouTube = None
    YouTubeTranscriptApi = None

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_OK = True
except Exception:
    SKLEARN_OK = False

try:
    import pandas as pd
except Exception:
    pd = None


# ------------------------------------------------------------------------------
# Configuration & Theming
# ------------------------------------------------------------------------------
st.set_page_config(page_title="🧠 Anki Deck Creator", page_icon="🧠", layout="wide")
st.markdown(
    """
<style>
:root {
  --card-border: #e5e7eb;
  --ink-1: #111827;
  --ink-2: #374151;
  --ink-3: #6b7280;
  --bg-soft: #fafbff;
  --brand: #4338ca;
  --brand-weak: #eef2ff;
}
.main > div { padding-top: 0 !important; }
.kpi { border:1px solid var(--card-border); border-radius:12px; padding:12px; text-align:center; }
.kpi h3 { margin:0; font-size:.9rem; color:var(--ink-2); }
.kpi .val { font-weight:700; font-size:1.1rem; }
.card-prev { border:1px solid var(--card-border); border-radius:12px; padding:12px; margin-bottom:10px; }
.step { padding:.6rem .8rem; border-radius:10px; border:1px solid var(--card-border); background:var(--bg-soft); margin-bottom:.55rem;}
.step .title { font-weight:600; color:var(--ink-1); }
.badge { display:inline-block; padding:.15rem .5rem; border-radius:999px; font-size:.75rem; background:var(--brand-weak); border:1px solid #e0e7ff; color:#3730a3; }
.source-chip { display:inline-block; margin:2px 6px 2px 0; padding:.2rem .5rem; background:#f3f4f6; border:1px solid #e5e7eb; border-radius:999px; font-size:.8rem; }
hr { border:none; border-top:1px solid var(--card-border); margin:1rem 0; }
</style>
""",
    unsafe_allow_html=True,
)

DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
TEXT_MODEL = DEFAULT_MODEL
MAX_AUDIO_FILES = 24
AUDIO_CHAR_LIMIT = 400
FALLBACK_PREVIEW = 3

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY", "")
if not OPENAI_API_KEY:
    st.error("Set **OPENAI_API_KEY** in Streamlit Secrets or environment.")
    st.stop()
client = OpenAI(api_key=OPENAI_API_KEY, timeout=60.0)

SYSTEM_PROMPT = """
Você é uma IA especialista em Design Instrucional e Ciência Cognitiva que atua como Criador de Baralhos Anki.
Trabalhe em 3 passos: (1) entender pedido e propor um plano, (2) gerar prévia variada e específica, (3) produzir o baralho final.

Regras para os cartões:
- Recordação ativa e conhecimento atômico (uma ideia por cartão).
- Perguntas específicas, resposta única, linguagem clara, concisa.
- Variedade além de Q&A: cloze (um único alvo), cenários/mini-casos, procedimentos/passos, contrastes/exceções.
- Para políticas/exames: prazos, responsáveis, thresholds, exceções, auditoria, conformidade.
- Para idiomas: vocabulário (collocations), gramática/padrões (contraste), listening/pronúncia (script natural/IPA quando útil).
- Quando houver materiais (payload.materials_digest), ancore cartões neles e preencha source_ref.file/page_or_time.
- NÃO dependa de materiais: se não houver, gere a partir do tópico e das metas.

Saída estritamente em JSON:
{
 "deck":{"title":"string","language":"string","level":"string","topic":"string","source_summary":"string","card_count_planned":number},
 "cards":[
  {"id":"string","type":"basic|reverse|cloze","front":"string","back":"string","hint":"string|null",
   "examples":[{"text":"string","translation":"string|null","notes":"string|null"}],
   "audio_script":"string|null","tags":["string"],"difficulty":"easy|medium|hard",
   "source_ref":{"file":"string|null","page_or_time":"string|null","span":"string|null"},
   "rationale":"string"}],
 "qa_report":{"duplicates_removed":number,"too_broad_removed":number,"notes":"string"}
}
Defina deck.card_count_planned = len(cards).
"""

ASSESSMENT_SYSTEM = """
Você é um especialista em Design Instrucional. Analise o pedido do usuário e (se houver) um digest dos materiais.
Proponha um plano sucinto para o baralho, em JSON:
{
 "summary":"string",
 "key_issues":["..."],
 "card_type_strategy":{"basic":int,"reverse":int,"cloze":int,"scenario":int,"procedure":int},
 "anchoring_plan":"string",
 "terminology_focus":["..."],
 "open_questions":["..."]
}
Os percentuais devem somar aproximadamente 100.
"""

# ------------------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------------------
def notify(msg: str, icon: str = "ℹ️") -> None:
    try: st.toast(msg, icon=icon)
    except Exception: st.info(msg)

def ss_init(k: str, v: Any) -> None:
    if k not in st.session_state: st.session_state[k] = v

for k, v in [
    ("topic", ""),
    ("deck_title", ""),
    ("goal", "Language: Vocabulary"),
    ("idioma", "en-US"),
    ("nivel", "B1"),
    ("tipos", ["basic", "reverse", "cloze"]),
    ("target_n", 36),
    ("max_qa_pct", 0.45),
    ("require_anchor", True),
    ("default_tag", ""),
    ("materials", []),
    ("urls_text", ""),
    ("youtube_urls_text", ""),
    ("digest", ""),
    ("chosen_sections_meta", []),
    ("domain_kws", []),
    ("assessment", None),
    ("preview_cards", []),
    ("approved_cards", []),
    ("tts_mode", "examples"),
    ("tts_cov", "sampled"),
]:
    ss_init(k, v)

# ------------------------------------------------------------------------------
# OpenAI robust JSON (no temperature)
# ------------------------------------------------------------------------------
def _extract_json_from_text(text: str) -> Dict[str, Any]:
    if not text: return {}
    try: return json.loads(text)
    except Exception: pass
    start, end = text.find("{"), text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try: return json.loads(text[start:end+1])
        except Exception: return {}
    return {}

def chat_json(messages: List[Dict[str, str]], model: str = TEXT_MODEL, max_tries: int = 4) -> Dict[str, Any]:
    last_err: Optional[Exception] = None
    for i in range(max_tries):
        try:
            resp = client.chat.completions.create(model=model, messages=messages, response_format={"type":"json_object"})
            content = resp.choices[0].message.content if resp and resp.choices else "{}"
            data = _extract_json_from_text(content)
            if data: return data
        except Exception as e:
            last_err = e
        try:
            lax_msgs = messages + [{"role":"user","content":"\nResponda APENAS em JSON válido."}]
            resp2 = client.chat.completions.create(model=model, messages=lax_msgs)
            content2 = resp2.choices[0].message.content if resp2 and resp2.choices else "{}"
            data2 = _extract_json_from_text(content2)
            if data2: return data2
        except Exception as e2:
            last_err = e2
        time.sleep(min(6, 1.5**i + random.random()))
    if last_err: st.error(f"Failed to reach the AI model: {last_err}")
    return {}

# ------------------------------------------------------------------------------
# Ingestion (files / URLs / YouTube)
# ------------------------------------------------------------------------------
def extract_pdf(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes); tmp.flush()
        reader = PdfReader(tmp.name)
        out: List[str] = []
        for i, p in enumerate(reader.pages, start=1):
            t = p.extract_text() or ""
            if t.strip(): out.append(f"[p.{i}]\n{t}")
        return "\n\n".join(out)

def extract_docx(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes); tmp.flush()
        doc = docxlib.Document(tmp.name)
        return "\n".join(p.text for p in doc.paragraphs)

def extract_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")

def ingest_files(uploaded_files) -> List[Dict[str, str]]:
    mats: List[Dict[str, str]] = []
    for f in uploaded_files or []:
        name, data = f.name, f.read()
        low = name.lower()
        if   low.endswith(".pdf"): text = extract_pdf(data)
        elif low.endswith(".docx"): text = extract_docx(data)
        elif low.endswith((".txt",".md",".markdown")): text = extract_txt(data)
        else: continue
        mats.append({"file": name, "content": text})
    return mats

def fetch_url_text(url: str, timeout: int = 25) -> str:
    if not requests: return ""
    try:
        r = requests.get(url, timeout=timeout, headers={"User-Agent":"anki-deck-creator/1.0"})
        r.raise_for_status()
        html = r.text
        if Document:
            try:
                doc = Document(html)
                content = doc.summary(html_partial=False)
                soup = BeautifulSoup(content, "html.parser")
                return soup.get_text("\n", strip=True)
            except Exception:
                pass
        if BeautifulSoup:
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script","style","noscript"]): tag.decompose()
            return soup.get_text("\n", strip=True)
        return html
    except Exception:
        return ""

def ingest_urls(urls: List[str]) -> List[Dict[str, str]]:
    mats: List[Dict[str, str]] = []
    for u in urls:
        u = u.strip()
        if not u: continue
        txt = fetch_url_text(u)
        if txt:
            name = re.sub(r"^https?://", "", u)[:80]
            mats.append({"file": name, "content": txt})
    return mats

def ingest_youtube_transcript(url: str) -> Optional[List[Dict[str, str]]]:
    if not YOUTUBE_DEPS_OK: return None
    try:
        yt = YouTube(url)
        video_id = yt.video_id
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
        for lang in ["en","pt","fr"]:
            try:
                transcript = transcript_list.find_transcript([lang])
                raw_text = transcript.fetch()
                full_transcript = " ".join([t['text'] for t in raw_text])
                title = yt.title
                return [{"file": f"youtube_{video_id}", "content": f"# {title}\n{full_transcript}"}]
            except Exception:
                continue
        return None
    except Exception as e:
        st.error(f"Error ingesting YouTube video: {e}")
        return None

# ------------------------------------------------------------------------------
# Lightweight RAG
# ------------------------------------------------------------------------------
def split_sections(text: str, file_name: str) -> List[Dict[str, str]]:
    sections: List[Dict[str, str]] = []
    parts = re.split(r"\n(?=#+\s)|\n(?=\[p\.\d+\])", text)
    for i, part in enumerate(parts):
        pt = part.strip()
        if not pt: continue
        m = re.match(r"^(#+\s.*)", pt)
        if m: title = m.group(1)
        elif pt.lower().startswith("[p."): title = pt.split("\n",1)[0]
        else: title = f"{file_name} — sec {i+1}"
        sections.append({"file": file_name, "title": title, "content": pt[:4000]})
    return sections or [{"file": file_name, "title": f"{file_name} — full", "content": text[:4000]}]

def simple_keyword_rank(query: str, sections: List[Dict[str, str]], top_k: int = 6) -> List[int]:
    q_toks = set(re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9]{3,}", query.lower()))
    scores: List[Tuple[int,int]] = []
    for i, s in enumerate(sections):
        txt = s["content"].lower()
        score = sum(1 for t in q_toks if t in txt)
        scores.append((score, i))
    scores.sort(reverse=True)
    return [i for _,i in scores[:top_k]]

def tfidf_rank(query: str, sections: List[Dict[str, str]], top_k: int = 6) -> List[int]:
    docs = [s["content"] for s in sections]
    try:
        vec = TfidfVectorizer(ngram_range=(1,2), min_df=1, max_df=0.9)
        X = vec.fit_transform(docs)
        qv = vec.transform([query])
        sims = cosine_similarity(qv, X).ravel()
        return sims.argsort()[::-1][:top_k].tolist()
    except Exception:
        return simple_keyword_rank(query, sections, top_k)

def rag_digest(materials: List[Dict[str, str]], topic: str, user_feedback: str = "", top_k: int = 6, max_chars: int = 12000) -> Tuple[str, List[Dict[str, str]]]:
    if not materials: return "", []
    sections: List[Dict[str, str]] = []
    for m in materials:
        sections.extend(split_sections(m["content"], m["file"]))
    query = f"{topic}\n{user_feedback}".strip()
    idxs = tfidf_rank(query, sections, top_k) if SKLEARN_OK else simple_keyword_rank(query, sections, top_k)
    chosen = [sections[i] for i in idxs]
    parts: List[str] = []
    for s in chosen:
        hdr = f"# {s['file']} — {s['title']}"
        parts.append(hdr + "\n" + s["content"])
    digest = "\n\n".join(parts)
    return digest[:max_chars], chosen

# ------------------------------------------------------------------------------
# Card validation & formatting
# ------------------------------------------------------------------------------
REQUIRED_CARD_FIELDS = {"type","front","back"}

def html_escape(s: Optional[str]) -> str:
    if not s: return ""
    return s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

def normalize_text_for_html(s: str, cloze: bool = False) -> str:
    if s is None: return ""
    s = str(s).strip()
    s = re.sub(r"^```+|\n```+$", "", s)
    s = html_escape(s) if not cloze else s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    return s.replace("\n","<br>")

def strip_html_to_plain(s: str) -> str:
    if not s: return ""
    s2 = re.sub(r"<[^>]+>", " ", s)
    s2 = (s2.replace("&nbsp;"," ").replace("&amp;","&").replace("&lt;","<").replace("&gt;",">"))
    return re.sub(r"\s+", " ", s2).strip()

def valid_card(c: dict) -> bool:
    if not isinstance(c, dict): return False
    typ = (c.get("type") or "").lower()
    if typ == "cloze":
        cloze_src = c.get("front","") or c.get("Text","")
        if "{{c" not in cloze_src: return False
        if len(strip_html_to_plain(c.get("back",""))) > 420: return False
        return True
    if not all(k in c and (c[k] is not None and len(str(c[k]).strip())>0) for k in REQUIRED_CARD_FIELDS): return False
    if len(strip_html_to_plain(c.get("back",""))) > 420: return False
    return True

def dedupe_cards(cards: list) -> list:
    def sig(c):
        t = (c.get("type") or "").lower().strip()
        f = (c.get("front") or c.get("Text") or "").strip().lower()
        b = (c.get("back") or c.get("BackExtra") or "").strip().lower()
        return hashlib.md5(f"{t}|{f[:160]}|{b[:160]}".encode()).hexdigest()
    seen, out = set(), []
    for c in cards or []:
        if not isinstance(c, dict): continue
        s = sig(c)
        if s in seen: continue
        seen.add(s); out.append(c)
    return out

def contains_domain_keyword(card: dict, kws: List[str]) -> bool:
    if not kws: return True
    blob = f"{card.get('front','')} | {card.get('back','')}".lower()
    return any(k in blob for k in kws[:40])

def anchored_to_source(card: dict) -> bool:
    src = card.get("source_ref") or {}
    return bool(src.get("file") or src.get("page_or_time"))

def kind_of(c: dict) -> str:
    t = (c.get("type") or "").lower()
    if t == "cloze": return "cloze"
    f = (c.get("front") or "").lower()
    if "?" in f: return "qa"
    if any(k in f for k in ["scenario:","cenário:","case:","caso:"]): return "scenario"
    if any(k in f for k in ["procedure","procedimento","step","passo","ordem"]): return "procedure"
    return "basic"

def enforce_variety(cards: List[dict], kws: List[str], max_qa_frac: float, require_anchor_effective: bool, seed_ids: set) -> List[dict]:
    filtered: List[dict] = []
    for c in cards:
        if require_anchor_effective and not anchored_to_source(c) and c.get("id") not in seed_ids: continue
        if not contains_domain_keyword(c, kws) and c.get("id") not in seed_ids: continue
        filtered.append(c)
    if not filtered: return cards
    total = max(1, len(filtered))
    qa_idx = [i for i,c in enumerate(filtered) if kind_of(c)=="qa" and c.get("id") not in seed_ids]
    max_qa = int(max_qa_frac * total)
    if len(qa_idx) > max_qa:
        drop = set(qa_idx[max_qa:])
        filtered = [c for i,c in enumerate(filtered) if i not in drop]
    return filtered

# ------------------------------------------------------------------------------
# LLM Pipelines
# ------------------------------------------------------------------------------
def _fallback_cards(payload: Dict[str, Any], n: int = FALLBACK_PREVIEW) -> List[Dict[str, Any]]:
    topic = (payload.get("topico") or payload.get("deck_title") or "Topic").strip()
    kws = payload.get("domain_keywords") or []
    src_file = None
    try:
        if st.session_state.get("materials"):
            src_file = st.session_state["materials"][0]["file"]
    except Exception:
        pass
    base = [
        {"type":"basic","front":f"What is {topic}?","back":f"{topic}: concise definition."},
        {"type":"cloze","front":f"{topic} involves {{c1::key ideas}} and practice.","back":"Reveal the hidden key idea."},
        {"type":"reverse","front":f"Term: {kws[0] if kws else topic}","back":f"Short explanation of {(kws[0] if kws else topic)}."},
    ]
    out=[]
    for i,c in enumerate(base[:max(1,n)]):
        cc=dict(c); cc.update({
            "id": f"fallback-{i}","hint":"","examples":[],"tags":["fallback","seed"],"difficulty":"medium",
            "source_ref":{"file":src_file or None,"page_or_time":None,"span":None}
        })
        out.append(cc)
    return out

def gerar_baralho(payload: Dict[str, Any]) -> Dict[str, Any]:
    messages = [
        {"role":"system","content":SYSTEM_PROMPT},
        {"role":"user","content":"Gere um baralho Anki a partir do payload a seguir (JSON):\n"+json.dumps(payload, ensure_ascii=False)}
    ]
    data = chat_json(messages, model=TEXT_MODEL) or {}
    cards = data.get("cards", [])
    if not isinstance(cards, list): cards = []

    # Normalize then validate
    normalized=[]
    for c in cards:
        if not isinstance(c, dict): continue
        typ = (c.get("type") or "basic").lower()
        if typ == "cloze":
            if not c.get("front") and c.get("Text"): c["front"] = str(c.get("Text",""))
            if not c.get("back"):
                be = str(c.get("BackExtra","")).strip()
                c["back"] = be if be else "(cloze answer on reveal)"
        c["front"] = str(c.get("front","")).strip()
        c["back"]  = str(c.get("back","")).strip()
        normalized.append(c)
    cards = normalized

    cards = [c for c in cards if valid_card(c)]
    for c in cards:
        typ = (c.get("type") or "basic").lower()
        if typ == "cloze":
            c["front"] = normalize_text_for_html(c.get("front",""), cloze=True)
            c["back"]  = normalize_text_for_html(c.get("back",""))
        else:
            c["front"] = normalize_text_for_html(c.get("front",""))
            c["back"]  = normalize_text_for_html(c.get("back",""))

    if not cards:
        cards = _fallback_cards(payload, n=max(1, int(payload.get("limite_cartoes", FALLBACK_PREVIEW))))

    deck = data.get("deck", {}) if isinstance(data.get("deck", {}), dict) else {}
    deck.setdefault("title", payload.get("deck_title","Anki Deck"))
    deck.setdefault("language", payload.get("idioma_alvo","en"))
    deck.setdefault("level", payload.get("nivel_proficiencia",""))
    deck.setdefault("topic", payload.get("topico",""))
    deck["card_count_planned"] = len(cards)
    return {"deck": deck, "cards": cards}

def generate_assessment(topic: str, goal: str, idioma: str, nivel: str, tipos: list,
                        digest: str, domain_kws: List[str],
                        max_qa_frac: float, require_anchor: bool) -> Dict[str, Any]:
    payload = {
        "topic": topic, "goal": goal, "language": idioma, "level": nivel,
        "allowed_types": tipos, "materials_digest_excerpt": (digest or "")[:6000],
        "domain_keywords": domain_kws[:30], "max_qa_pct": max_qa_frac, "require_anchor": require_anchor
    }
    messages = [
        {"role":"system","content":ASSESSMENT_SYSTEM},
        {"role":"user","content":"Analise o pedido/materiais e proponha um plano de baralho (JSON):\n"+json.dumps(payload, ensure_ascii=False)}
    ]
    data = chat_json(messages, model=TEXT_MODEL) or {}
    data.setdefault("summary","")
    data.setdefault("key_issues",[])
    data.setdefault("card_type_strategy",{"basic":35,"reverse":10,"cloze":35,"scenario":10,"procedure":10})
    data.setdefault("anchoring_plan","Reference file + page/section when available. If no materials, rely on topic-only generation.")
    data.setdefault("terminology_focus", domain_kws[:12])
    data.setdefault("open_questions",[])
    return data

# ------------------------------------------------------------------------------
# TTS & Genanki
# ------------------------------------------------------------------------------
QUESTION_WORDS = {
    "fr": ["qui","que","quoi","quel","quelle","quels","quelles","où","quand","comment","pourquoi","combien"],
    "pt": ["o que","qual","quais","quando","onde","como","por que","por quê","quem","quanto","quantos","quantas"],
    "en": ["what","which","when","where","how","why","who","whom","whose"],
}
def _lang_prefix(lang: str) -> str:
    l = (lang or "en").lower()
    for k in ["fr","pt","en","es","de","it"]:
        if l.startswith(k): return k
    return "en"
def looks_like_question(text: str, lang: str) -> bool:
    if not text: return False
    t = text.strip()
    if "?" in t: return True
    lp = _lang_prefix(lang)
    w = re.sub(r'^[\-\•\–\s]+', '', t.lower())
    return any(w.startswith(q + " ") or w == q for q in QUESTION_WORDS.get(lp,[]))
def orient_q_a(card: Dict[str, Any], lang: str) -> Tuple[str, str]:
    f = (card.get("front") or card.get("Text") or "").strip()
    b = (card.get("back") or card.get("BackExtra") or "").strip()
    fq, bq = looks_like_question(f, lang), looks_like_question(b, lang)
    if fq and not bq: return f, b
    if bq and not fq: return b, f
    if len(f) > len(b): return b, f
    return f, b

def slugify(text: str, maxlen: int = 64) -> str:
    t = unidecode(text or "").lower()
    t = re.sub(r"[^\w\s-]", "", t)
    t = re.sub(r"[\s-]+", "-", t).strip("-")
    return t[:maxlen] or f"slug-{int(time.time())}"

def examples_to_html(examples: Optional[List[Dict[str, str]]]) -> str:
    if not examples: return ""
    items: List[str] = []
    for ex in examples:
        if not isinstance(ex, dict): continue
        line = html_escape(ex.get("text",""))
        tr = ex.get("translation"); nt = ex.get("notes")
        if tr: line += f' <span class="tr">({html_escape(tr)})</span>'
        if nt: line += f' <span class="nt">[{html_escape(nt)}]</span>'
        items.append(f"<li>{line}</li>")
    return "<ul class='examples-list'>" + "".join(items) + "</ul>"

def map_lang_for_gtts(language_code: str) -> str:
    langs = tts_langs()
    if not language_code: return "en"
    lc = language_code.lower()
    prefs = {"fr":["fr"],"pt":["pt","pt-br"],"en":["en","en-us","en-gb"],"es":["es"],"de":["de"],"it":["it"],"zh":["zh","zh-cn","cmn-cn","zh-tw"],"ja":["ja"]}
    for k, choices in prefs.items():
        if lc.startswith(k):
            for code in choices:
                if code in langs: return code
    for cand in (lc, lc.replace("_","-")):
        if cand in langs: return cand
    return "en"

def synth_audio(text: str, lang_code: str) -> Optional[bytes]:
    t = (text or "").strip()
    if not t or len(t) > AUDIO_CHAR_LIMIT: return None
    try:
        fp = io.BytesIO()
        gTTS(text=t, lang=map_lang_for_gtts(lang_code)).write_to_fp(fp)
        return fp.getvalue()
    except Exception:
        return None

def _clean_tag(tag) -> str:
    t = str(tag or "").strip().lower()
    t = re.sub(r"\s+", "-", t); t = re.sub(r"[^a-z0-9_\-:]", "", t)
    return t[:40]

def sanitize_tags(tags) -> list:
    if not isinstance(tags, list): return []
    out, seen = [], set()
    for t in tags:
        ct = _clean_tag(t)
        if ct and ct not in seen:
            out.append(ct); seen.add(ct)
        if len(out) >= 12: break
    return out

def stable_model_id(name: str, version: int = 3) -> int:
    h = hashlib.sha1(f"{name}-v{version}".encode("utf-8")).hexdigest()
    return int(h[:10], 16)

COMMON_CSS = """
.card { font-family: -apple-system, Segoe UI, Roboto, Arial; font-size: 20px; text-align: left; color: #222; background: #fff; }
.front { font-size: 1.05em; line-height: 1.45; }
.back { line-height: 1.5; }
.hint { margin-top: 8px; font-size: 0.95em; color: #666; }
.examples-list { padding-left: 18px; }
.tr { color: #444; font-style: italic; }
.nt { color: #666; }
.extra { margin-top: 10px; font-size: 0.9em; color: #444; }
.audio { margin-top: 8px; }
hr { margin: 14px 0; }
"""
MODEL_BASIC = genanki.Model(
    stable_model_id("Anki-Chat Basic", version=3), "Anki-Chat Basic (v3)",
    fields=[{"name":"Front"},{"name":"Back"},{"name":"Hint"},{"name":"Examples"},{"name":"Audio"},{"name":"Extra"}],
    templates=[{"name":"Card 1","qfmt":"<div class='front'>{{Front}}</div><div class='hint'>{{Hint}}</div>","afmt":"{{FrontSide}}<hr id='answer'><div class='back'>{{Back}}</div><div class='examples'>{{Examples}}</div><div class='audio'>{{Audio}}</div><div class='extra'>{{Extra}}</div>"}],
    css=COMMON_CSS
)
MODEL_REVERSE = genanki.Model(
    stable_model_id("Anki-Chat Reverse", version=3), "Anki-Chat Reverse (v3)",
    fields=[{"name":"Front"},{"name":"Back"},{"name":"Hint"},{"name":"Examples"},{"name":"Audio"},{"name":"Extra"}],
    templates=[
        {"name":"Forward","qfmt":"<div class='front'>{{Front}}</div><div class='hint'>{{Hint}}</div>","afmt":"{{FrontSide}}<hr id='answer'><div class='back'>{{Back}}</div><div class='examples'>{{Examples}}</div><div class='audio'>{{Audio}}</div><div class='extra'>{{Extra}}</div>"},
        {"name":"Reverse","qfmt":"<div class='front'>{{Back}}</div><div class='hint'>{{Hint}}</div>","afmt":"{{FrontSide}}<hr id='answer'><div class='back'>{{Front}}</div><div class='examples'>{{Examples}}</div><div class='audio'>{{Audio}}</div><div class='extra'>{{Extra}}</div>"},
    ],
    css=COMMON_CSS
)
MODEL_CLOZE = genanki.Model(
    stable_model_id("Anki-Chat Cloze", version=3), "Anki-Chat Cloze (v3)",
    fields=[{"name":"Text"},{"name":"BackExtra"},{"name":"Hint"},{"name":"Examples"},{"name":"Audio"},{"name":"Extra"}],
    templates=[{"name":"Cloze","qfmt":"{{cloze:Text}}<div class='hint'>{{Hint}}</div>","afmt":"{{cloze:Text}}<hr id='answer'><div class='back'>{{BackExtra}}</div><div class='examples'>{{Examples}}</div><div class='audio'>{{Audio}}</div><div class='extra'>{{Extra}}</div>"}],
    css=COMMON_CSS, model_type=genanki.Model.CLOZE
)

def build_deck_id(title: str) -> int:
    h = hashlib.sha1(title.encode("utf-8")).hexdigest()
    return int(h[:10], 16)

def choose_tts_text(card: Dict[str, Any], policy: str, lang: str, front_raw: str, back_raw: str) -> Optional[str]:
    if policy == "none": return None
    audio_script = (card.get("audio_script") or "").strip()
    ex = ""
    if isinstance(card.get("examples"), list) and card["examples"]:
        ex = (card["examples"][0].get("text") or "").strip()
    p = policy.lower()
    if   p == "all":      candidates = [audio_script, ex, back_raw, front_raw]
    elif p == "answers":  candidates = [audio_script, back_raw, ex, front_raw]
    elif p == "examples": candidates = [ex, audio_script, back_raw, front_raw]
    else:                 candidates = [audio_script, back_raw, ex, front_raw]
    for c in candidates:
        t = strip_html_to_plain(c)
        if t: return t
    return None

def build_apkg_bytes(deck_json: Dict[str, Any], tts_policy: str = "examples", tts_coverage: str = "sampled", default_tag: str = "") -> bytes:
    meta = deck_json.get("deck", {})
    cards = deck_json.get("cards", [])
    title = meta.get("title","Anki Deck")
    lang = meta.get("language","en")
    def anchored(c): return 1 if anchored_to_source(c) else 0
    def diff_rank(c): return {"easy":0,"medium":1,"hard":2}.get((c.get("difficulty") or "").lower(), 1)
    cards = sorted(cards, key=lambda c: (-anchored(c), c.get("type",""), diff_rank(c), c.get("id","")))
    deck = genanki.Deck(build_deck_id(title), title)

    media_files: List[str] = []
    tmpdir = tempfile.mkdtemp(prefix="anki_media_")
    idxs_for_audio = set()
    if tts_policy!= "none" and len(cards) > 0:
        if tts_coverage.lower().startswith("all"): idxs_for_audio = set(range(len(cards)))
        else: idxs_for_audio = set(random.sample(range(len(cards)), min(MAX_AUDIO_FILES, len(cards))))
    audio_map: Dict[int,str] = {}

    if idxs_for_audio:
        from concurrent.futures import ThreadPoolExecutor, as_completed
        def synth_one(idx):
            c = cards[idx]; fr, br = orient_q_a(c, lang)
            txt = choose_tts_text(c, tts_policy, lang, fr, br)
            if not txt: return idx, "", ""
            b = synth_audio(txt, lang)
            if not b: return idx, "", ""
            p = os.path.join(tmpdir, f"tts_{int(time.time()*1000)}_{idx}.mp3")
            with open(p, "wb") as f: f.write(b)
            return idx, f"[sound:{os.path.basename(p)}]", p
        with ThreadPoolExecutor(max_workers=5) as ex:
            futs = [ex.submit(synth_one, i) for i in idxs_for_audio]
            with st.spinner("Generating audio..."):
                for fut in as_completed(futs):
                    try:
                        i, field, path = fut.result()
                        if field:
                            audio_map[i] = field
                            media_files.append(path)
                    except Exception:
                        pass

    def add_note(c: Dict[str, Any], index: int):
        ctype = (c.get("type") or "basic").lower()
        hint = html_escape(c.get("hint") or "")
        examples_html = examples_to_html(c.get("examples"))
        src = c.get("source_ref") or {}
        extra_bits: List[str] = []
        if (src.get("file") or src.get("page_or_time")):
            extra_bits.append(f"<div><b>Source:</b> {html_escape(src.get('file') or '')} {html_escape(src.get('page_or_time') or '')}</div>")
        extra = "".join(extra_bits)
        audio_field = audio_map.get(index, "")
        front_raw, back_raw = orient_q_a(c, lang)
        front = normalize_text_for_html(front_raw, cloze=(ctype=="cloze"))
        back  = normalize_text_for_html(back_raw)
        tags = sanitize_tags(c.get("tags", []))
        if default_tag: tags = sanitize_tags(tags + [default_tag])
        guid = genanki.guid_for(front, back)
        if ctype == "cloze":
            note = genanki.Note(model=MODEL_CLOZE, fields=[front, back, hint, examples_html, audio_field, extra], tags=tags, guid=guid)
        elif ctype == "reverse":
            note = genanki.Note(model=MODEL_REVERSE, fields=[front, back, hint, examples_html, audio_field, extra], tags=tags, guid=guid)
        else:
            note = genanki.Note(model=MODEL_BASIC, fields=[front, back, hint, examples_html, audio_field, extra], tags=tags, guid=guid)
        deck.add_note(note)

    prog = st.progress(0, text="Building Anki deck...")
    for i, c in enumerate(cards):
        if isinstance(c, dict): add_note(c, i)
        prog.progress((i+1)/max(1,len(cards)), text=f"Adding card {i+1} of {len(cards)}...")
    prog.empty()

    pkg = genanki.Package(deck)
    if media_files: pkg.media_files = media_files
    with tempfile.NamedTemporaryFile(delete=False, suffix=".apkg") as tmp:
        pkg.write_to_file(tmp.name)
        with open(tmp.name, "rb") as f:
            return f.read()

# ------------------------------------------------------------------------------
# Stats & Rendering
# ------------------------------------------------------------------------------
def deck_stats(cards, domain_kws=None):
    kinds = [kind_of(c) for c in cards]
    total = len(cards) or 1
    cover = 0
    if domain_kws:
        blob = " ".join(strip_html_to_plain((c.get("front",""))+" "+(c.get("back",""))) for c in cards).lower()
        cover = 100*sum(1 for k in domain_kws[:20] if k in blob)/max(1, min(20, len(domain_kws)))
    anch = 100*sum(1 for c in cards if anchored_to_source(c))/total
    return {
      "total": total,
      "qa_pct": 100*sum(1 for k in kinds if k=="qa")/total,
      "cloze_pct": 100*sum(1 for k in kinds if k=="cloze")/total,
      "scenario_pct": 100*sum(1 for k in kinds if k=="scenario")/total,
      "procedure_pct": 100*sum(1 for k in kinds if k=="procedure")/total,
      "avg_back_len": sum(len(strip_html_to_plain(c.get("back",""))) for c in cards)/total,
      "anchored_pct": anch,
      "coverage_pct": cover
    }

def render_card_preview(c: Dict[str, Any], lang: str) -> str:
    typ = (c.get("type") or "basic").title()
    fr, br = orient_q_a(c, lang)
    front = normalize_text_for_html(fr, cloze=(typ.lower()=="cloze"))
    back  = normalize_text_for_html(br)
    hint = html_escape(c.get("hint") or "")
    src = c.get("source_ref") or {}
    src_str = " • ".join([x for x in [src.get("file"), src.get("page_or_time")] if x])
    exs = examples_to_html(c.get("examples"))
    parts = [
        "<div class='card-prev'>",
        f"<div style='font-weight:600;opacity:.8;margin-bottom:6px;'>{typ}</div>",
        f"<div><b>Q:</b> {front}</div>",
        f"<div style='margin-top:6px;'><b>A:</b> {back}</div>",
    ]
    if hint: parts.append(f"<div style='margin-top:6px;color:#666;'><b>Hint:</b> {hint}</div>")
    if exs:  parts.append(f"<div style='margin-top:6px;'><b>Examples:</b> {exs}</div>")
    if src_str: parts.append(f"<div style='margin-top:6px;color:#444;'><b>Source:</b> {html_escape(src_str)}</div>")
    parts.append("</div>")
    return "".join(parts)

# ------------------------------------------------------------------------------
# Payload builders & batching to reach target
# ------------------------------------------------------------------------------
def require_anchor_effective() -> bool:
    return bool(st.session_state.get("require_anchor", True) and (st.session_state.get("materials") or st.session_state.get("digest")))
def make_payload(n_cards: int, seed_cards=None):
    return {
        "deck_title": st.session_state.get("deck_title") or "Anki Deck",
        "idioma_alvo": st.session_state.get("idioma", "en"),
        "nivel_proficiencia": st.session_state.get("nivel", ""),
        "topico": st.session_state.get("topic") or "(no topic)",
        "limite_cartoes": n_cards,
        "tipos_permitidos": st.session_state.get("tipos", ["basic","reverse","cloze"]),
        "politica_voz": f"tts={st.session_state.get('tts_mode','examples')}",
        "materials_digest": st.session_state.get("digest", ""),
        "extra_policy": "usage_tip",
        "goal": st.session_state.get("goal", "Language: Vocabulary"),
        "max_qa_pct": st.session_state.get("max_qa_pct", 0.45),
        "require_anchor": require_anchor_effective(),
        "domain_keywords": (st.session_state.get("domain_kws") or [])[:30],
        "user_feedback": "",
        "seed_cards": seed_cards or [],
        "assessment_plan": st.session_state.get("assessment") or {},
        "assessment_approved": bool(st.session_state.get("assessment") is not None),
    }

def build_sample(payload: Dict[str, Any]) -> Dict[str, Any]:
    data = gerar_baralho(payload)
    cards = dedupe_cards(data.get("cards", []))
    seed_ids = {c.get("id") for c in payload.get("seed_cards", []) if c.get("id")}
    cards = enforce_variety(cards, st.session_state.get("domain_kws", []), payload["max_qa_pct"], payload["require_anchor"], seed_ids)
    data["cards"] = (cards or [])[: payload["limite_cartoes"]]
    data["deck"]["card_count_planned"] = len(data["cards"])
    return data

def build_to_target(target_n: int, seed_cards: List[dict], max_rounds: int = 6, batch_cap: int = 24, progress=None) -> List[dict]:
    cards = dedupe_cards(seed_cards or [])
    seen_ids = {hashlib.md5(((c.get('id') or '') + strip_html_to_plain(c.get('front',''))[:120] + strip_html_to_plain(c.get('back',''))[:120]).encode()).hexdigest() for c in cards}
    rounds = 0
    no_gain_rounds = 0

    # primeira atualização de UI
    if progress is not None:
        progress.progress(min(1.0, len(cards)/max(1, target_n)), text=f"Gerando cartões… {len(cards)}/{target_n}")

    while len(cards) < target_n and rounds < max_rounds:
        remain = target_n - len(cards)
        ask_n = min(batch_cap, remain)
        seed_hint = cards[-min(20, len(cards)):]
        payload = make_payload(ask_n, seed_cards=seed_hint)

        data = gerar_baralho(payload)
        more = dedupe_cards(data.get("cards", []) or [])
        more = enforce_variety(
            more, st.session_state.get("domain_kws", []),
            payload["max_qa_pct"], payload["require_anchor"],
            {c.get("id") for c in seed_hint if c.get("id")}
        ) or []

        new_cards = []
        for c in more:
            sig = hashlib.md5(((c.get('id') or '') + strip_html_to_plain(c.get('front',''))[:120] + strip_html_to_plain(c.get('back',''))[:120]).encode()).hexdigest()
            if sig in seen_ids: 
                continue
            seen_ids.add(sig)
            new_cards.append(c)

        if not new_cards:
            # relaxa restrições se travar
            payload["require_anchor"] = False
            payload["max_qa_pct"] = max(payload["max_qa_pct"], 0.9)
            data = gerar_baralho(payload)
            more = dedupe_cards(data.get("cards", []) or [])
            new_cards = []
            for c in more:
                sig = hashlib.md5(((c.get('id') or '') + strip_html_to_plain(c.get('front',''))[:120] + strip_html_to_plain(c.get('back',''))[:120]).encode()).hexdigest()
                if sig in seen_ids:
                    continue
                seen_ids.add(sig)
                new_cards.append(c)

        prev_len = len(cards)
        cards = dedupe_cards(cards + new_cards)

        rounds += 1
        if len(cards) == prev_len:
            no_gain_rounds += 1
        else:
            no_gain_rounds = 0

        # atualização a cada rodada
        if progress is not None:
            pct = min(1.0, len(cards)/max(1, target_n))
            progress.progress(pct, text=f"Gerando cartões… {len(cards)}/{target_n} (lote {rounds}/{max_rounds})")

        if no_gain_rounds >= 2:
            break

    if progress is not None:
        pct = min(1.0, len(cards)/max(1, target_n))
        progress.progress(pct, text=f"Concluído: {len(cards)}/{target_n}")
        time.sleep(0.3)
        progress.empty()

    return cards[:target_n]

# ------------------------------------------------------------------------------
# Header
# ------------------------------------------------------------------------------
h_l, h_r = st.columns([3, 2])
with h_l:
    st.title("🧠 Anki Deck Creator")
    st.caption("Manual workflow • Plan → Preview → Build high-quality Anki decks from topics & sources.")
with h_r:
    m1, m2, m3 = st.columns(3)
    m1.metric("Target", st.session_state["target_n"])
    m2.metric("Max Q&A %", f"{int(st.session_state['max_qa_pct']*100)}%")
    # Model override
    new_model = st.text_input("OpenAI model", value=TEXT_MODEL, help="Override via OPENAI_MODEL env var.")
    if new_model and new_model != TEXT_MODEL:
        TEXT_MODEL = new_model
        notify(f"Model set to {TEXT_MODEL}.", "✅")

top_cols = st.columns([1,1,2,1])
with top_cols[0]:
    if st.button("🔄 Reset session"):
        # Simple reset of key state
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()
with top_cols[1]:
    st.caption("Tip: Add sources to increase anchoring & variety.")
with top_cols[2]:
    st.caption("1) Generate plan → 2) Preview & edit → 3) Build & download")
with top_cols[3]:
    pass

st.divider()

# ------------------------------------------------------------------------------
# Layout: Left (Wizard) / Right (Sources & Build Settings)
# ------------------------------------------------------------------------------
left, right = st.columns([2, 1])

with right:
    st.subheader("📎 Sources (optional)")
    with st.expander("Add / Manage Sources", expanded=False):
        upl = st.file_uploader("Upload files", type=["pdf","docx","txt","md","markdown"], accept_multiple_files=True)
        url_lines = st.text_area("URLs (one per line)", st.session_state["urls_text"], height=80)
        youtube_urls = st.text_area("YouTube URLs (one per line)", st.session_state["youtube_urls_text"], height=80)
        c1, c2 = st.columns(2)
        if c1.button("Add sources"):
            with st.spinner("Processing sources..."):
                mats = ingest_files(upl) if upl else []
                if url_lines.strip():
                    st.session_state["urls_text"] = url_lines
                    urls = [u.strip() for u in url_lines.splitlines() if u.strip()]
                    mats += ingest_urls(urls)
                if youtube_urls.strip() and YOUTUBE_DEPS_OK:
                    st.session_state["youtube_urls_text"] = youtube_urls
                    yt_urls = [u.strip() for u in youtube_urls.splitlines() if u.strip()]
                    for u in yt_urls:
                        yt_mat = ingest_youtube_transcript(u) or []
                        mats += yt_mat
                if mats:
                    st.session_state["materials"].extend(mats)
                    notify(f"Added {len(mats)} new sources!", "✅")
                else:
                    notify("No new sources added.", "ℹ️")
        if c2.button("Clear sources"):
            st.session_state["materials"] = []
            st.session_state["digest"] = ""
            st.session_state["chosen_sections_meta"] = []
            notify("Sources cleared.", "🗑️")

    if st.session_state["materials"]:
        st.markdown("**Current sources**")
        for m in st.session_state["materials"]:
            st.markdown(f"<span class='source-chip'>{m['file']}</span>", unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("⚙️ Build Settings")
    st.session_state["tts_mode"] = st.selectbox("TTS", ["none","answers","examples","all"], index=["none","answers","examples","all"].index(st.session_state["tts_mode"]))
    st.session_state["tts_cov"]  = st.selectbox("TTS coverage", ["sampled","all"], index=["sampled","all"].index(st.session_state["tts_cov"]))
    st.session_state["default_tag"] = st.text_input("Default tag", st.session_state["default_tag"] or slugify(st.session_state.get("topic","")))
    st.session_state["require_anchor"] = st.checkbox("Require anchoring (auto-disabled if no materials)", value=st.session_state["require_anchor"])
    st.session_state["max_qa_pct"] = st.slider("Max % Q&A", 10, 90, int(st.session_state["max_qa_pct"]*100), 5)/100.0
    st.session_state["target_n"] = st.slider("Target cards", 4, 200, st.session_state["target_n"], 2)

with left:
    # Stepper
    st.markdown("<div class='step'><span class='title'>1) Plan</span> <span class='badge'>Assessment</span><br><span class='help'>Set topic/goal/language. (Optional) Add sources to anchor.</span></div>", unsafe_allow_html=True)
    plan_cols = st.columns(2)
    with plan_cols[0]:
        st.session_state["topic"] = st.text_input("Topic", st.session_state["topic"])
        st.session_state["deck_title"] = st.text_input("Deck title", st.session_state["deck_title"] or st.session_state["topic"][:48])
        st.session_state["idioma"] = st.text_input("Language code", st.session_state["idioma"])
    with plan_cols[1]:
        st.session_state["nivel"] = st.text_input("Level", st.session_state["nivel"])
        st.session_state["goal"] = st.selectbox(
            "Goal",
            ["General learning","Org policy mastery","Exam prep","Language: Vocabulary","Language: Grammar & Patterns"],
            index=["General learning","Org policy mastery","Exam prep","Language: Vocabulary","Language: Grammar & Patterns"].index(st.session_state["goal"])
        )
        st.session_state["tipos"] = st.multiselect("Allowed types", ["basic","reverse","cloze"], default=st.session_state["tipos"])

    # Build digest / keywords if materials exist
    if st.session_state["materials"]:
        digest, chosen = rag_digest(st.session_state["materials"], st.session_state["topic"], "", top_k=6)
        st.session_state["digest"] = digest
        st.session_state["chosen_sections_meta"] = chosen
        toks = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9\-_/\.]{4,}", digest.lower())
        st.session_state["domain_kws"] = list(dict.fromkeys(toks))[:40]
    else:
        st.session_state["digest"] = ""
        toks = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9\-_/\.]{4,}", (st.session_state["topic"] or "").lower())
        st.session_state["domain_kws"] = list(dict.fromkeys(toks))[:20]

    plan_actions = st.columns([1,1,3])
    gen_plan = plan_actions[0].button("🧭 Generate plan")
    if gen_plan:
        if not st.session_state["topic"].strip():
            st.warning("Add a topic first.")
        else:
            a = generate_assessment(
                st.session_state["topic"], st.session_state["goal"], st.session_state["idioma"], st.session_state["nivel"],
                st.session_state["tipos"], st.session_state["digest"], st.session_state["domain_kws"],
                st.session_state["max_qa_pct"], st.session_state["require_anchor"]
            )
            st.session_state["assessment"] = a
            st.success("Plan generated below.")

    if st.session_state["assessment"]:
        st.json(st.session_state["assessment"], expanded=False)

    st.markdown("<div class='step'><span class='title'>2) Preview</span> <span class='badge'>Sample cards</span><br><span class='help'>Review & edit a small sample. Approved rows become seed cards.</span></div>", unsafe_allow_html=True)
    prev_cols = st.columns([1,1,3])
    prev_size = prev_cols[0].selectbox("Preview size", [1,2,3,4,5], index=2)
    do_preview = prev_cols[1].button("🔎 Preview")

    if do_preview:
        if not st.session_state["topic"].strip():
            st.warning("Add a topic first.")
        else:
            payload = make_payload(prev_size)
            data = build_sample(payload)
            st.session_state["preview_cards"] = data.get("cards", [])
            stats = deck_stats(st.session_state["preview_cards"], st.session_state["domain_kws"])
            st.info(
                f"Preview ready ({len(st.session_state['preview_cards'])} cards). "
                f"Mix — Q&A {stats['qa_pct']:.0f}%, Cloze {stats['cloze_pct']:.0f}%, "
                f"Scenario {stats['scenario_pct']:.0f}%, Procedure {stats['procedure_pct']:.0f}%. "
                f"Anchored {stats['anchored_pct']:.0f}%."
            )

    if st.session_state["preview_cards"]:
        deck_lang = st.session_state["idioma"]
        cols = st.columns(2)
        for i, c in enumerate(st.session_state["preview_cards"]):
            html = render_card_preview(c, deck_lang)
            (cols[i % 2]).markdown(html, unsafe_allow_html=True)

        st.caption("Approve/edit the sample (kept rows become seed cards for the final deck).")
        if pd:
            def cards_to_df(cards):
                rows=[]
                for c in cards:
                    rows.append({
                        "keep": True,
                        "type": c.get("type","basic"),
                        "front": strip_html_to_plain(c.get("front") or c.get("Text","")),
                        "back":  strip_html_to_plain(c.get("back") or c.get("BackExtra","")),
                        "hint": c.get("hint",""),
                        "tags": ",".join(sanitize_tags(c.get("tags", []))),
                    })
                return pd.DataFrame(rows)
            def df_to_cards(df):
                out=[]
                for _,r in df.iterrows():
                    if not r.get("keep", False): continue
                    tags = sanitize_tags([t.strip() for t in str(r.get("tags","")).split(",") if t.strip()])
                    cid = f"seed-{hashlib.md5((str(r['type'])+str(r['front'])).encode()).hexdigest()[:8]}"
                    out.append({
                        "id": cid, "type": r["type"],
                        "front": normalize_text_for_html(r["front"], cloze=(r["type"]=="cloze")),
                        "back":  normalize_text_for_html(r["back"]),
                        "hint":  html_escape(r.get("hint","")),
                        "examples": [], "tags": tags,
                        "source_ref": {}, "rationale": ""
                    })
                return out
            df = cards_to_df(st.session_state["preview_cards"])
            edited = st.data_editor(df, use_container_width=True, num_rows="fixed", key="preview_editor")
            st.session_state["approved_cards"] = df_to_cards(edited)
        else:
            keep_opts = [st.checkbox(f"Keep card {i+1}", value=True, key=f"keep_{i}") for i in range(len(st.session_state["preview_cards"]))]
            st.session_state["approved_cards"] = [c for i,c in enumerate(st.session_state["preview_cards"]) if keep_opts[i]]

    st.markdown("<div class='step'><span class='title'>3) Build</span> <span class='badge'>.apkg export</span><br><span class='help'>Generate full deck with optional TTS and download.</span></div>", unsafe_allow_html=True)
    do_build = st.button("🏗️ Build deck")
    if do_build:
        if not st.session_state["topic"].strip():
            st.warning("Add a topic first.")
        else:
            with st.spinner("Building your deck…"):
                seeds = st.session_state["approved_cards"] or st.session_state["preview_cards"] or []
                cards = build_to_target(st.session_state["target_n"], seeds)
                deck_json = {
                    "deck": {
                        "title": st.session_state["deck_title"] or st.session_state["topic"][:48] or "Anki Deck",
                        "language": st.session_state["idioma"],
                        "level": st.session_state["nivel"],
                        "topic": st.session_state["topic"],
                        "source_summary": ""
                    },
                    "cards": cards
                }
                apkg = build_apkg_bytes(
                    deck_json,
                    tts_policy=st.session_state["tts_mode"],
                    tts_coverage=("all" if st.session_state["tts_cov"].startswith("all") else "sampled"),
                    default_tag=(st.session_state["default_tag"] or slugify(st.session_state["topic"]))
                )
                fname = f"{slugify(deck_json['deck']['title'])}_{int(time.time())}.apkg"
                st.success(f"Deck built with {len(deck_json['cards'])} cards!")
                st.download_button("⬇️ Download .apkg", data=apkg, file_name=fname, mime="application/octet-stream")
                stats = deck_stats(deck_json["cards"], st.session_state["domain_kws"])
                k1,k2,k3,k4,k5,k6 = st.columns(6)
                for k, title, val in [
                    (k1,"Cards", f"{stats['total']}"),
                    (k2,"Q&A %", f"{stats['qa_pct']:.0f}%"),
                    (k3,"Cloze %", f"{stats['cloze_pct']:.0f}%"),
                    (k4,"Scenario %", f"{stats['scenario_pct']:.0f}%"),
                    (k5,"Anchored %", f"{stats['anchored_pct']:.0f}%"),
                    (k6,"Coverage", f"{stats['coverage_pct']:.0f}%"),
                ]:
                    with k:
                        st.markdown(f"<div class='kpi'><h3>{title}</h3><div class='val'>{val}</div></div>", unsafe_allow_html=True)
